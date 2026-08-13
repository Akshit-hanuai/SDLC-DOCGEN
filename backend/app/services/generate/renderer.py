from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.schemas.template import SectionSchema, TemplateSchema


def render_document(
    schema: TemplateSchema,
    project_name: str,
    content: dict,
    out_path: str | Path,
    classification: str = "UNCLASSIFIED",
) -> str:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _title_page(doc, schema, project_name, content.get("header", {}), classification)
    doc.add_page_break()

    sections = content.get("sections", {})
    for section in schema.sections:
        _render_section(doc, section, sections.get(section.id))

    doc.save(str(out_path))
    return str(out_path)


def _title_page(doc: Document, schema: TemplateSchema, project_name: str, header: dict, classification: str):
    doc.add_paragraph(f"DOCUMENT CLASSIFICATION: {classification}")
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(schema.name)
    run.bold = True
    run.font.size = Pt(24)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Project: {project_name}").font.size = Pt(14)
    doc.add_paragraph()
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for field_def in schema.header.document_control:
        value = header.get(field_def.field, field_def.default or "")
        row = table.add_row()
        row.cells[0].text = field_def.field.replace("_", " ").title()
        row.cells[1].text = str(value)


def _render_section(doc: Document, section: SectionSchema, content):
    doc.add_heading(f"{section.id} {section.title}", level=1)
    if not content:
        doc.add_paragraph("(not generated)")
        return

    if section.type == "requirements":
        _render_requirement_section(doc, section, content)
        return
    if section.type == "traceability_matrix":
        _render_matrix(doc, content)
        return

    for field in section.fields:
        doc.add_heading(field.title, level=2)
        value = content.get(field.id)
        if field.type == "table":
            _render_table(doc, field.columns or [], value)
        elif field.type == "list":
            if isinstance(value, list):
                for item in value:
                    text = item.get("text") if isinstance(item, dict) else str(item)
                    doc.add_paragraph(text, style="List Bullet")
            else:
                doc.add_paragraph(str(value or ""))
        else:
            doc.add_paragraph(str(value or ""))


def _render_requirement_section(doc: Document, section: SectionSchema, content):
    rows = content.get("requirements", [])
    if not isinstance(rows, list):
        rows = []
    columns = list((section.output.columns if section.output else None) or ["requirement_id", "requirement"])
    if section.annexure:
        doc.add_heading("Annexure", level=1)
    _render_table(doc, columns, rows)
    gaps = content.get("gaps", [])
    if gaps:
        doc.add_heading("Gap Analysis", level=2)
        for gap in gaps:
            doc.add_paragraph(gap, style="List Bullet")


def _render_matrix(doc: Document, content):
    rows = content.get("rows", [])
    if isinstance(rows, list) and rows:
        _render_table(doc, list(rows[0].keys()), rows)
    uncovered = content.get("uncovered", [])
    if uncovered:
        doc.add_heading("Uncovered Requirements", level=2)
        for item in uncovered:
            doc.add_paragraph(item, style="List Bullet")


def _render_table(doc: Document, columns: list[str], rows):
    rows = rows or []
    if not rows:
        doc.add_paragraph("(empty)")
        return
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for i, column in enumerate(columns):
        table.rows[0].cells[i].text = column.replace("_", " ").title()
    for row in rows:
        if isinstance(row, dict):
            values = [str(row.get(c, "")) for c in columns]
        else:
            values = [str(v) for v in row]
        cells = table.add_row().cells
        for i, value in enumerate(values):
            if i < len(cells):
                cells[i].text = value
