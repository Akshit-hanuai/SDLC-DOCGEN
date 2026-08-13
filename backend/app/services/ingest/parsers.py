import re
from pathlib import Path

from app.services.ingest.models import Block, ParsedDocument

HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")


def parse_file(path: str, doc_type: str) -> ParsedDocument:
    suffix = Path(path).suffix.lower()
    if suffix in (".docx", ".doc"):
        return _parse_docx(path, doc_type)
    if suffix == ".pdf":
        return _parse_pdf(path, doc_type)
    return _parse_text(path, doc_type)


def _blocks_from_lines(lines: list[tuple[str, int]], filename: str, doc_type: str) -> ParsedDocument:
    blocks: list[Block] = []
    current = Block(heading="", level=1, text="")
    for text, level in lines:
        stripped = text.strip()
        if not stripped:
            continue
        if level > 0:
            if current.text.strip() or current.heading:
                blocks.append(current)
            current = Block(heading=stripped, level=level, text="")
        else:
            current.text += stripped + "\n"
    if current.text.strip() or current.heading:
        blocks.append(current)
    full_text = "\n".join(b.heading + "\n" + b.text for b in blocks)
    return ParsedDocument(filename=filename, doc_type=doc_type, text=full_text, blocks=blocks)


def _parse_text(path: str, doc_type: str) -> ParsedDocument:
    content = Path(path).read_text(encoding="utf-8", errors="replace")
    lines: list[tuple[str, int]] = []
    for raw in content.splitlines():
        line = raw.rstrip()
        if not line.strip():
            lines.append(("", 0))
            continue
        m = HEADING_RE.match(line.strip())
        if m:
            lines.append((m.group(2), len(m.group(1))))
            continue
        if line.strip().isupper() and len(line.strip()) > 4:
            lines.append((line.strip(), 2))
            continue
        lines.append((line, 0))
    return _blocks_from_lines(lines, Path(path).name, doc_type)


def _parse_docx(path: str, doc_type: str) -> ParsedDocument:
    from docx import Document

    doc = Document(path)
    lines: list[tuple[str, int]] = []
    for para in doc.paragraphs:
        # para.style can be None for paragraphs with missing/corrupt style refs
        style_name = ""
        if para.style is not None and para.style.name:
            style_name = para.style.name.lower()
        text = para.text.rstrip()
        if not text.strip():
            continue
        if "heading" in style_name:
            digits = "".join(ch for ch in style_name if ch.isdigit())
            level = int(digits) if digits else 1
            lines.append((text, level))
        else:
            lines.append((text, 0))
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append((" | ".join(cells), 0))
    return _blocks_from_lines(lines, Path(path).name, doc_type)


def _parse_pdf(path: str, doc_type: str) -> ParsedDocument:
    import pdfplumber

    lines: list[tuple[str, int]] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                for line in (page.extract_text() or "").splitlines():
                    stripped = line.strip()
                    if not stripped:
                        continue
                    if _looks_like_heading(stripped):
                        lines.append((stripped, 2))
                    else:
                        lines.append((stripped, 0))
    except Exception:
        import fitz

        doc = fitz.open(path)
        for page in doc:
            for line in page.get_text().splitlines():
                stripped = line.strip()
                if not stripped:
                    continue
                if _looks_like_heading(stripped):
                    lines.append((stripped, 2))
                else:
                    lines.append((stripped, 0))
    return _blocks_from_lines(lines, Path(path).name, doc_type)


def _looks_like_heading(line: str) -> bool:
    if len(line) > 80:
        return False
    if HEADING_RE.match(line):
        return True
    if line.isupper() and len(line) > 4:
        return True
    return bool(re.match(r"^\d+(\.\d+)*\s+[A-Z]", line))
