"""Generate sample SDLC input fixtures (SysRS, IRS, MoM) as DOCX files.

Usage: python scripts/gen_fixtures.py
Outputs: sample_data/fixtures/sysrs.docx, irs.docx, mom.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent.parent / "sample_data" / "fixtures"
OUT.mkdir(parents=True, exist_ok=True)

SYSRS_REQUIREMENTS = [
    ("REQ-0001", "functional", "The system shall acquire telemetry data from at least 3 sensor channels simultaneously."),
    ("REQ-0002", "non_functional", "The system shall have a mean time between failures (MTBF) of not less than 2000 hours."),
    ("REQ-0003", "functional", "The system shall timestamp every acquired sample with a system time reference accurate to +/- 1 ms."),
    ("REQ-0004", "functional", "The system shall buffer at least 10000 samples per channel before flow control is applied. This refines REQ-0003 by requiring bounded buffering for time-tagged data."),
    ("REQ-0005", "functional", "The system shall expose acquired samples through a REST interface to the ground segment."),
    ("REQ-0006", "functional", "The system shall record an engineering log with a minimum depth of 6 months of mission operations."),
    ("REQ-0007", "non_functional", "The system shall sustain a sample ingestion rate of at least 100 Hz per channel."),
    ("REQ-0008", "non_functional", "The system shall restrict access to acquisition controls to authenticated operators only."),
    ("REQ-0009", "non_functional", "The system shall complete a software restart in less than 30 seconds after a fault."),
    ("REQ-0010", "constraint", "The system shall operate within a 28 V DC power envelope not exceeding 45 W."),
]

IRS_REQUIREMENTS = [
    ("IR-0101", "REQ-0005", "The ground segment interface shall transmit telemetry sample envelopes as defined in api.proto (message TelemetrySample)."),
    ("IR-0102", "REQ-0005", "The acquisition command interface shall accept a StartCommand message to begin acquisition."),
    ("IR-0103", "REQ-0002", "The health reporting interface shall expose an MTBF estimate as defined in api.proto (message HealthReport)."),
    ("IR-0104", "REQ-0001", "The sensor interface shall provide a raw sample stream at the configured sample rate."),
]


def _base_doc(title: str) -> Document:
    doc = Document()
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc


def _reqs_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, header in enumerate(["ID", "Type", "Requirement"]):
        table.rows[0].cells[i].text = header
    for req_id, req_type, text in rows:
        cells = table.add_row().cells
        cells[0].text = req_id
        cells[1].text = req_type
        cells[2].text = text


def gen_sysrs() -> Path:
    doc = _base_doc("System Requirements Specification (SysRS)")
    doc.add_paragraph("Revision 3.1 — baseline for pilot-project-alpha.")
    doc.add_heading("1. Scope", level=1)
    doc.add_paragraph("This SysRS defines the system-level requirements for the telemetry acquisition subsystem.")
    doc.add_heading("2. Requirements", level=1)
    doc.add_heading("2.1 Functional Requirements", level=2)
    _reqs_table(doc, [r for r in SYSRS_REQUIREMENTS if r[1] == "functional"])
    doc.add_heading("2.2 Non-Functional Requirements", level=2)
    _reqs_table(doc, [r for r in SYSRS_REQUIREMENTS if r[1] != "functional"])
    path = OUT / "SysRS_v3.1.docx"
    doc.save(path)
    return path


def gen_irs() -> Path:
    doc = _base_doc("Interface Requirements Specification (IRS)")
    doc.add_paragraph("Revision 2.0 — interface requirements for the ground segment and sensor interfaces.")
    doc.add_heading("3. Interface Requirements", level=1)
    rows = [(req_id, "interface", text) for req_id, _, text in IRS_REQUIREMENTS]
    _reqs_table(doc, rows)
    path = OUT / "IRS_v2.0.docx"
    doc.save(path)
    return path


def gen_mom() -> Path:
    doc = _base_doc("Minutes of Meeting — Design Review 04")
    doc.add_paragraph("Review of the telemetry acquisition subsystem preliminary design. Rev 1.0")
    doc.add_heading("1. Decisions", level=1)
    doc.add_paragraph("Decision: The acquisition subsystem shall use a time-tagged sample buffer for all channels.")
    doc.add_paragraph("Decision: The REST interface in REQ-0005 shall be versioned from day one.")
    doc.add_heading("2. Action Items", level=1)
    doc.add_paragraph("Action item: Update SysRS section 2 to clarify the REQ-0001 sampling rate. Owner: Dr. Sharma. Due: 2026-09-01.")
    doc.add_paragraph("Action item: Confirm REQ-0008 authentication mechanism with the security panel. Owner: Lt. Rao. Due: 2026-09-15.")
    doc.add_heading("3. Requirement Changes", level=1)
    doc.add_paragraph("Requirement change: REQ-0006 is modified to require structured JSON logging output.")
    doc.add_paragraph("Requirement change: A new requirement REQ-0011 is added to require automatic fault recovery within 30 seconds, in addition to REQ-0009.")
    doc.add_heading("4. Test Outcomes", level=1)
    doc.add_paragraph("REQ-0001: PASS — all 3 channels acquired during the hardware-in-loop demonstration.")
    doc.add_paragraph("REQ-0005: FAIL — sample retrieval endpoint did not meet latency budget during load test.")
    doc.add_paragraph("REQ-0002: PASS — MTBF estimate recorded above 2000 hours on the reference platform.")
    path = OUT / "MoM_DesignReview04.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    for name, fn in [("sysrs", gen_sysrs), ("irs", gen_irs), ("mom", gen_mom)]:
        path = fn()
        print(f"generated {path}")
